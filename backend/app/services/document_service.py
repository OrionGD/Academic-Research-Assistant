"""
Document Service - Orchestrates document processing pipeline.
Manages file uploads, text extraction, chunking, metadata extraction, and storage coordination.
"""

import os
import shutil
import logging
import tempfile
from pathlib import Path
from typing import Optional, List, Dict, Any, BinaryIO, Tuple
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed

from sqlalchemy.orm import Session
import magic
from PyPDF2 import PdfReader
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter, TextSplitter
import tiktoken

from app.models.database import Document, DocumentChunk, User
from app.models.schemas import DocumentCreate, DocumentUpdate
from app.core.config import settings
from app.utils.file_utils import FileUtils
from app.utils.text_utils import TextUtils

logger = logging.getLogger(__name__)


class DocumentService:
    """Service for document processing operations"""
    
    SUPPORTED_FILE_TYPES = {
        'application/pdf': 'pdf',
        'application/msword': 'doc',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'text/plain': 'txt',
        'text/markdown': 'md',
        'text/csv': 'csv',
        'application/json': 'json',
        'text/html': 'html',
    }
    
    MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB
    CHUNK_SIZE = 1000
    CHUNK_OVERLAP = 200
    
    def __init__(self, db: Session):
        self.db = db
        self.file_utils = FileUtils()
        self.text_utils = TextUtils()
        
    async def upload_document(
        self,
        user_id: int,
        file_obj: BinaryIO,
        filename: str,
        title: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Document:
        """
        Upload and process a document.
        
        Args:
            user_id: ID of the user uploading the document
            file_obj: File-like object containing document data
            filename: Original filename
            title: Optional document title (defaults to filename)
            metadata: Additional document metadata
            
        Returns:
            Document object
        """
        # Validate user exists
        user = self.db.query(User).filter(User.id == user_id, User.is_active.is_(True)).first()
        if not user:
            raise ValueError(f"User with ID {user_id} not found or inactive")
        
        # Read file content for validation
        file_content = file_obj.read()
        file_size = len(file_content)
        
        # Validate file size
        if file_size > self.MAX_FILE_SIZE:
            raise ValueError(f"File size {file_size} exceeds maximum {self.MAX_FILE_SIZE}")
        
        # Determine file type
        file_type = self._detect_file_type(file_content, filename)
        if not file_type:
            raise ValueError(f"Unsupported file type: {filename}")
        
        # Create document record
        title = title or Path(filename).stem
        document_data = DocumentCreate(
            title=title,
            filename=filename,
            file_size=file_size,
            file_type=file_type,
            metadata=metadata or {}
        )
        
        # Generate safe filename
        safe_filename = self.file_utils.generate_safe_filename(filename)
        upload_dir = self._get_user_upload_dir(user_id)
        file_path = upload_dir / safe_filename
        
        # Save file
        upload_dir.mkdir(parents=True, exist_ok=True)
        with open(file_path, 'wb') as f:
            f.write(file_content)
        
        # Create database record
        document = Document(
            user_id=user_id,
            title=document_data.title,
            filename=document_data.filename,
            file_path=str(file_path),
            file_size=document_data.file_size,
            file_type=document_data.file_type,
            metadata=document_data.metadata,
            processed=False,
            processing_status='pending'
        )
        
        self.db.add(document)
        self.db.commit()
        self.db.refresh(document)
        
        # Start async processing
        self._process_document_async(document.id)
        
        logger.info(f"Document {document.id} uploaded by user {user_id}")
        return document
    
    def _detect_file_type(self, content: bytes, filename: str) -> Optional[str]:
        """Detect file type using magic numbers and extension"""
        # Try magic detection first
        mime = magic.from_buffer(content[:2048], mime=True)
        if mime in self.SUPPORTED_FILE_TYPES:
            return self.SUPPORTED_FILE_TYPES[mime]
        
        # Fallback to extension
        ext = Path(filename).suffix.lower()[1:]  # Remove leading dot
        if ext in ['pdf']:
            return 'pdf'
        elif ext in ['doc', 'docx']:
            return 'docx' if ext == 'docx' else 'doc'
        elif ext in ['txt', 'md', 'csv', 'json', 'html']:
            return ext
        
        return None
    
    def _get_user_upload_dir(self, user_id: int) -> Path:
        """Get upload directory for user"""
        upload_base = Path(settings.UPLOAD_DIR)
        return upload_base / str(user_id)
    
    def _process_document_async(self, document_id: int):
        """Start document processing in background thread"""
        from app.core.celery_app import celery_app
        
        # Send task to Celery worker
        celery_app.send_task(
            'app.worker.process_document',
            args=[document_id],
            queue='document_processing'
        )
    
    def process_document(self, document_id: int) -> Dict[str, Any]:
        """
        Process document: extract text, chunk, generate embeddings.
        
        Args:
            document_id: ID of document to process
            
        Returns:
            Processing results
        """
        document = self.db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise ValueError(f"Document {document_id} not found")
        
        try:
            # Update processing status
            document.processing_status = 'processing'
            self.db.commit()
            
            # Extract text
            logger.info(f"Extracting text from document {document_id}")
            text_content = self._extract_text(document.file_path, document.file_type)
            
            if not text_content.strip():
                raise ValueError("No text content extracted from document")
            
            # Extract metadata
            logger.info(f"Extracting metadata from document {document_id}")
            extracted_metadata = self._extract_metadata(document, text_content)
            document.metadata.update(extracted_metadata)
            
            # Split into chunks
            logger.info(f"Splitting document {document_id} into chunks")
            chunks = self._split_text(text_content, document.metadata)
            
            # Create chunk records
            logger.info(f"Creating {len(chunks)} chunk records for document {document_id}")
            chunk_objects = []
            
            for i, chunk in enumerate(chunks):
                chunk_obj = DocumentChunk(
                    document_id=document_id,
                    chunk_index=i,
                    content=chunk['text'],
                    chunk_size=len(chunk['text']),
                    token_count=self.text_utils.count_tokens(chunk['text']),
                    metadata=chunk.get('metadata', {})
                )
                chunk_objects.append(chunk_obj)
                self.db.add(chunk_obj)
            
            # Update document status
            document.processed = True
            document.processing_status = 'completed'
            self.db.commit()
            
            logger.info(f"Successfully processed document {document_id} with {len(chunks)} chunks")
            
            return {
                'document_id': document_id,
                'chunks_count': len(chunks),
                'status': 'completed'
            }
            
        except Exception as e:
            logger.error(f"Error processing document {document_id}: {str(e)}")
            document.processing_status = 'failed'
            self.db.commit()
            raise
    
    def _extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text content from different file types"""
        if file_type == 'pdf':
            return self._extract_text_from_pdf(file_path)
        elif file_type in ['doc', 'docx']:
            return self._extract_text_from_docx(file_path)
        elif file_type == 'txt':
            return self._extract_text_from_txt(file_path)
        elif file_type == 'md':
            return self._extract_text_from_txt(file_path)
        elif file_type == 'csv':
            return self._extract_text_from_csv(file_path)
        elif file_type == 'json':
            return self._extract_text_from_json(file_path)
        elif file_type == 'html':
            return self._extract_text_from_html(file_path)
        else:
            raise ValueError(f"Unsupported file type for text extraction: {file_type}")
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text_parts = []
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                for page_num, page in enumerate(reader.pages, 1):
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(f"Page {page_num}:\n{text}")
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise
    
    def _extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            raise ValueError("Unable to decode text file with common encodings")
    
    def _extract_text_from_csv(self, file_path: str) -> str:
        """Extract text from CSV file"""
        import csv
        try:
            text_parts = []
            with open(file_path, 'r', encoding='utf-8') as file:
                reader = csv.reader(file)
                for i, row in enumerate(reader, 1):
                    row_text = " | ".join([cell for cell in row if cell.strip()])
                    if row_text:
                        text_parts.append(f"Row {i}: {row_text}")
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting text from CSV: {str(e)}")
            raise
    
    def _extract_text_from_json(self, file_path: str) -> str:
        """Extract text from JSON file"""
        import json
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
                # Convert JSON to readable text
                return self._json_to_text(data)
        except Exception as e:
            logger.error(f"Error extracting text from JSON: {str(e)}")
            raise
    
    def _extract_text_from_html(self, file_path: str) -> str:
        """Extract text from HTML file"""
        from bs4 import BeautifulSoup
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file, 'html.parser')
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                return soup.get_text(separator='\n', strip=True)
        except Exception as e:
            logger.error(f"Error extracting text from HTML: {str(e)}")
            raise
    
    def _json_to_text(self, data: Any, indent: int = 0) -> str:
        """Convert JSON data to readable text"""
        if isinstance(data, dict):
            text_parts = []
            for key, value in data.items():
                value_text = self._json_to_text(value, indent + 2)
                text_parts.append(f"{' ' * indent}{key}: {value_text}")
            return "\n".join(text_parts)
        elif isinstance(data, list):
            text_parts = []
            for i, item in enumerate(data):
                item_text = self._json_to_text(item, indent + 2)
                text_parts.append(f"{' ' * indent}[{i}]: {item_text}")
            return "\n".join(text_parts)
        else:
            return str(data)
    
    def _extract_metadata(self, document: Document, text_content: str) -> Dict[str, Any]:
        """Extract metadata from document and content"""
        metadata = {
            'characters_count': len(text_content),
            'words_count': len(text_content.split()),
            'lines_count': text_content.count('\n') + 1,
            'processing_timestamp': datetime.utcnow().isoformat(),
        }
        
        # Try to extract title if not in metadata
        if 'title' not in document.metadata:
            # Extract first line as potential title
            first_line = text_content.split('\n')[0].strip()[:200]
            if first_line:
                metadata['extracted_title'] = first_line
        
        # Extract language
        from langdetect import detect, LangDetectException
        try:
            sample_text = text_content[:500]
            language = detect(sample_text)
            metadata['detected_language'] = language
        except LangDetectException:
            metadata['detected_language'] = 'unknown'
        
        return metadata
    
    def _split_text(self, text: str, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Split text into chunks with metadata.
        
        Args:
            text: Text to split
            metadata: Document metadata
            
        Returns:
            List of chunks with text and metadata
        """
        # Create text splitter
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.CHUNK_SIZE,
            chunk_overlap=self.CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # Split text
        text_chunks = splitter.split_text(text)
        
        # Create chunk objects with metadata
        chunks = []
        for i, chunk_text in enumerate(text_chunks):
            chunk_metadata = {
                'chunk_index': i,
                'total_chunks': len(text_chunks),
                'source_document_metadata': metadata.copy()
            }
            chunks.append({
                'text': chunk_text,
                'metadata': chunk_metadata
            })
        
        return chunks
    
    def get_user_documents(
        self,
        user_id: int,
        skip: int = 0,
        limit: int = 100,
        processed_only: bool = False
    ) -> Tuple[List[Document], int]:
        """
        Get documents for a user with pagination.
        
        Args:
            user_id: User ID
            skip: Number of documents to skip
            limit: Maximum number of documents to return
            processed_only: Only return processed documents
            
        Returns:
            Tuple of (documents list, total count)
        """
        query = self.db.query(Document).filter(Document.user_id == user_id)
        
        if processed_only:
            query = query.filter(Document.processed.is_(True))
        
        total = query.count()
        documents = query.order_by(Document.upload_date.desc()).offset(skip).limit(limit).all()
        
        return documents, total
    
    def get_document_by_id(self, document_id: int, user_id: Optional[int] = None) -> Optional[Document]:
        """
        Get document by ID, optionally checking user ownership.
        
        Args:
            document_id: Document ID
            user_id: Optional user ID for ownership check
            
        Returns:
            Document if found and accessible, None otherwise
        """
        query = self.db.query(Document)
        
        if user_id is not None:
            query = query.filter(Document.user_id == user_id)
        
        return query.filter(Document.id == document_id).first()
    
    def update_document(
        self,
        document_id: int,
        user_id: int,
        update_data: DocumentUpdate
    ) -> Optional[Document]:
        """
        Update document metadata.
        
        Args:
            document_id: Document ID
            user_id: User ID for ownership verification
            update_data: Update data
            
        Returns:
            Updated document or None if not found/accessible
        """
        document = self.get_document_by_id(document_id, user_id)
        if not document:
            return None
        
        # Update fields
        if update_data.title is not None:
            document.title = update_data.title
        
        if update_data.metadata is not None:
            document.metadata.update(update_data.metadata)
        
        document.updated_at = datetime.utcnow()
        
        self.db.commit()
        self.db.refresh(document)
        
        return document
    
    def delete_document(self, document_id: int, user_id: int) -> bool:
        """
        Delete document and associated chunks.
        
        Args:
            document_id: Document ID
            user_id: User ID for ownership verification
            
        Returns:
            True if deleted, False if not found/accessible
        """
        document = self.get_document_by_id(document_id, user_id)
        if not document:
            return False
        
        try:
            # Delete file
            if os.path.exists(document.file_path):
                os.remove(document.file_path)
            
            # Delete from database (cascade will delete chunks)
            self.db.delete(document)
            self.db.commit()
            
            logger.info(f"Deleted document {document_id} for user {user_id}")
            return True
            
        except Exception as e:
            logger.error(f"Error deleting document {document_id}: {str(e)}")
            self.db.rollback()
            return False
    
    def get_document_chunks(
        self,
        document_id: int,
        user_id: Optional[int] = None,
        skip: int = 0,
        limit: int = 100
    ) -> Tuple[List[DocumentChunk], int]:
        """
        Get chunks for a document.
        
        Args:
            document_id: Document ID
            user_id: Optional user ID for ownership verification
            skip: Number of chunks to skip
            limit: Maximum number of chunks to return
            
        Returns:
            Tuple of (chunks list, total count)
        """
        # Verify document access
        document = self.get_document_by_id(document_id, user_id)
        if not document:
            return [], 0
        
        # Get chunks
        query = self.db.query(DocumentChunk).filter(DocumentChunk.document_id == document_id)
        total = query.count()
        chunks = query.order_by(DocumentChunk.chunk_index).offset(skip).limit(limit).all()
        
        return chunks, total